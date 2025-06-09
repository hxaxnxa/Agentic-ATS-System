import PyPDF2
import docx
import logging
import re
from io import BytesIO
import xml.etree.ElementTree as ET
import zipfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def parse_document(file, filename):
    """
    Parse document (PDF or DOCX) with enhanced resume-specific text extraction.
    Extracts text from paragraphs, tables, headers, footers, and handles common resume formats.
    Returns structured text with identified resume sections or raises an exception on failure.
    """
    file.seek(0)
    if filename.endswith(".docx"):
        try:
            doc = docx.Document(file)
            text = []
            current_section = "General"

            # Common resume section headers (case-insensitive)
            section_headers = [
                r'education\b', r'experience\b', r'work experience\b', r'skills\b',
                r'projects\b', r'certifications\b', r'awards\b', r'publications\b',
                r'objective\b', r'summary\b', r'professional summary\b'
            ]
            section_regex = re.compile('|'.join(section_headers), re.IGNORECASE)

            # Extract text from paragraphs
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if not para_text:
                    continue

                # Identify section headers
                if section_regex.search(para_text):
                    current_section = para_text
                    text.append(f"\n[SECTION: {current_section}]\n")
                else:
                    # Handle bullet points and clean text
                    cleaned_text = re.sub(r'^\s*[-•*]\s+', '', para_text)  # Remove bullet symbols
                    if cleaned_text:
                        text.append(f"{cleaned_text}\n")

            # Extract text from tables (common for skills or experience in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            para_text = paragraph.text.strip()
                            if para_text:
                                # Remove bullet symbols in table cells
                                cleaned_text = re.sub(r'^\s*[-•*]\s+', '', para_text)
                                cell_text += cleaned_text + " "
                        if cell_text.strip():
                            row_text.append(cell_text.strip())
                    if row_text:
                        text.append(" | ".join(row_text) + "\n")

            # Extract text from headers and footers (e.g., candidate name, contact info)
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text.append(f"[HEADER] {para.text.strip()}\n")
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text.append(f"[FOOTER] {para.text.strip()}\n")

            # Clean up the text
            full_text = "".join(text)
            full_text = re.sub(r'\n+', '\n', full_text)  # Remove multiple newlines
            full_text = re.sub(r'[ \t]+', ' ', full_text)  # Normalize spaces

            if full_text.strip():
                logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
                return full_text.strip()
            else:
                # Fall back to advanced DOCX extraction for complex layouts
                logger.warning("Standard DOCX parsing yielded no text, attempting advanced extraction")
                advanced_text = extract_text_from_docx_advanced(file)
                if advanced_text and advanced_text.strip():
                    logger.info(f"Advanced extraction succeeded, extracted {len(advanced_text)} characters")
                    return advanced_text.strip()
                raise ValueError("No text content found in DOCX")

        except Exception as e:
            logger.error(f"python-docx parsing failed: {e}")
            # Attempt advanced extraction as a fallback
            advanced_text = extract_text_from_docx_advanced(file)
            if advanced_text and advanced_text.strip():
                logger.info(f"Advanced extraction succeeded, extracted {len(advanced_text)} characters")
                return advanced_text.strip()
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    elif filename.endswith(".pdf"):
        try:
            reader = PyPDF2.PdfReader(file)
            text = []
            current_section = "General"

            # Common resume section headers (case-insensitive)
            section_headers = [
                r'education\b', r'experience\b', r'work experience\b', r'skills\b',
                r'projects\b', r'certifications\b', r'awards\b', r'publications\b',
                r'objective\b', r'summary\b', r'professional summary\b'
            ]
            section_regex = re.compile('|'.join(section_headers), re.IGNORECASE)

            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    if not page_text.strip():
                        continue

                    # Split page text into lines for better section detection
                    lines = page_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue

                        # Identify section headers
                        if section_regex.search(line):
                            current_section = line
                            text.append(f"\n[SECTION: {current_section}]\n")
                        else:
                            # Remove bullet symbols and clean text
                            cleaned_line = re.sub(r'^\s*[-•*]\s+', '', line)
                            if cleaned_line:
                                text.append(f"{cleaned_line}\n")

                    # Normalize spaces within the page text
                    page_text = re.sub(r'[ \t]+', ' ', page_text)
                    page_text = re.sub(r'\n\s*\n', '\n', page_text)

                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue

            # Clean up the text
            full_text = "".join(text)
            full_text = re.sub(r'\n+', '\n', full_text)  # Remove multiple newlines
            full_text = re.sub(r'[ \t]+', ' ', full_text)  # Normalize spaces

            if full_text.strip():
                logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
                return full_text.strip()
            else:
                raise ValueError("No text content found in PDF")

        except Exception as e:
            logger.error(f"PyPDF2 parsing failed: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    else:
        raise ValueError("Unsupported file format")

def extract_text_from_docx_advanced(file):
    """
    Advanced DOCX text extraction with better handling of complex resume layouts.
    Extracts text from document.xml and includes potential resume section detection.
    """
    try:
        # Read DOCX as ZIP file
        with zipfile.ZipFile(file, 'r') as docx_zip:
            # Extract document.xml
            doc_xml = docx_zip.read('word/document.xml')
            root = ET.fromstring(doc_xml)

            text = []
            current_section = "General"

            # Common resume section headers (case-insensitive)
            section_headers = [
                r'education\b', r'experience\b', r'work experience\b', r'skills\b',
                r'projects\b', r'certifications\b', r'awards\b', r'publications\b',
                r'objective\b', r'summary\b', r'professional summary\b'
            ]
            section_regex = re.compile('|'.join(section_headers), re.IGNORECASE)

            # Extract all text nodes
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_content = elem.text.strip()
                    # Identify section headers
                    if section_regex.search(text_content):
                        current_section = text_content
                        text.append(f"\n[SECTION: {current_section}]\n")
                    else:
                        # Remove bullet symbols
                        cleaned_text = re.sub(r'^\s*[-•*]\s+', '', text_content)
                        if cleaned_text:
                            text.append(f"{cleaned_text}\n")

            # Clean up the text
            full_text = "".join(text)
            full_text = re.sub(r'\n+', '\n', full_text)  # Remove multiple newlines
            full_text = re.sub(r'[ \t]+', ' ', full_text)  # Normalize spaces

            return full_text.strip()

    except Exception as e:
        logger.error(f"Advanced DOCX extraction failed: {e}")
        return None